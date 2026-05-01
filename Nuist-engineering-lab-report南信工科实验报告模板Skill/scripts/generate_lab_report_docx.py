from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def u(text: str) -> str:
    return text.encode("ascii").decode("unicode_escape")


FONT_SONG = u("\\u5b8b\\u4f53")
FONT_HEI = u("\\u9ed1\\u4f53")
DEFAULT_COVER_TITLE = u("\\uff08\\uff09\\u5b9e\\u9a8c\\u62a5\\u544a")
DEFAULT_EXPERIMENT_TITLE = u("\\u5b9e\\u9a8c\\u62a5\\u544a")
BODY_HEADINGS = {
    u("\\u5b9e\\u9a8c\\u76ee\\u7684"): u("\\u4e00\\u3001\\u5b9e\\u9a8c\\u76ee\\u7684"),
    u("\\u5b9e\\u9a8c\\u8fc7\\u7a0b"): u("\\u4e8c\\u3001\\u5b9e\\u9a8c\\u8fc7\\u7a0b"),
    u("\\u5b9e\\u9a8c\\u7ed3\\u679c"): u("\\u4e09\\u3001\\u5b9e\\u9a8c\\u7ed3\\u679c"),
    u("\\u5b9e\\u9a8c\\u7ed3\\u679c\\u5206\\u6790"): u("\\u56db\\u3001\\u5b9e\\u9a8c\\u7ed3\\u679c\\u5206\\u6790"),
}


def strip_inline_code_marks(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.replace("`", "")


def today_cn() -> str:
    today = date.today()
    digits = {str(i): c for i, c in enumerate(u("\\u3007\\u4e00\\u4e8c\\u4e09\\u56db\\u4e94\\u516d\\u4e03\\u516b\\u4e5d"))}
    year = "".join(digits[ch] for ch in str(today.year))
    return f"{year}{u('\\u5e74')} {today.month} {u('\\u6708')} {today.day} {u('\\u65e5')}"


def set_run_font(run, chinese_font: str, western_font: str, size_pt: float, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = western_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), chinese_font)
    rfonts.set(qn("w:ascii"), western_font)
    rfonts.set(qn("w:hAnsi"), western_font)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm: float | None = 0.74):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = None if first_line_cm is None else Cm(first_line_cm)


def replace_paragraph_text(paragraph, text: str, chinese_font: str, size_pt: float, bold: bool):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, chinese_font, "Times New Roman", size_pt, bold)


def set_experiment_line(paragraph, experiment_title: str):
    for run in paragraph.runs:
        run.text = ""
    run1 = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run1.text = experiment_title
    set_run_font(run1, FONT_SONG, "Times New Roman", 18, True)
    run1.font.underline = True
    run2 = paragraph.add_run("        ")
    set_run_font(run2, FONT_SONG, "Times New Roman", 18, False)
    run2.font.underline = True


def add_body_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph(p)
    run = p.add_run(strip_inline_code_marks(text))
    set_run_font(run, FONT_SONG, "Times New Roman", 10.5)


def add_heading(document: Document, text: str, level: int):
    p = document.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0)
    run = p.add_run(strip_inline_code_marks(text))
    set_run_font(run, FONT_HEI, "Times New Roman", 14 if level == 2 else 10.5, True)


def add_list_item(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0)
    run = p.add_run(strip_inline_code_marks(text))
    set_run_font(run, FONT_SONG, "Times New Roman", 10.5)


def add_image_with_caption(document: Document, image_path: Path, caption: str):
    p = document.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0)
    p.add_run().add_picture(str(image_path), width=Cm(14.5))
    cap = document.add_paragraph()
    set_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0)
    run = cap.add_run(strip_inline_code_marks(caption))
    set_run_font(run, FONT_SONG, "Times New Roman", 9)


def parse_image(line: str):
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
    if not match:
        return None
    return match.group(1), Path(match.group(2).replace("/", "\\"))


def guess_experiment_title(markdown_text: str, markdown_path: Path) -> str:
    for line in markdown_text.splitlines():
        text = line.strip()
        if text.startswith("# ") and text[2:].strip():
            return text[2:].strip()
    stem = markdown_path.parent.name.strip()
    return stem if stem else DEFAULT_EXPERIMENT_TITLE


def prepare_cover_from_template(document: Document, experiment_title: str):
    nonempty = [p for p in document.paragraphs if p.text.strip()]
    if len(nonempty) >= 2:
        replace_paragraph_text(nonempty[0], DEFAULT_COVER_TITLE, FONT_HEI, 18, True)
        set_experiment_line(nonempty[1], experiment_title)
        replace_paragraph_text(nonempty[-1], today_cn(), FONT_SONG, 18, False)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def build_body(document: Document, markdown_text: str):
    paragraph_lines: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            add_body_paragraph(document, " ".join(paragraph_lines))
            paragraph_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if line.startswith("## "):
            flush_paragraph()
            add_heading(document, BODY_HEADINGS.get(line[3:], line[3:]), 2)
            continue
        if line.startswith("### "):
            flush_paragraph()
            add_heading(document, line[4:], 3)
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            add_list_item(document, line)
            continue
        parsed = parse_image(line)
        if parsed:
            flush_paragraph()
            caption, image_path = parsed
            if image_path.exists():
                add_image_with_caption(document, image_path, caption)
            continue
        paragraph_lines.append(line)
    flush_paragraph()


def resolve_template(markdown_path: Path, template_arg: str | None) -> Path:
    if template_arg:
        return Path(template_arg)
    sibling_template = markdown_path.parent / u("\\u5b9e\\u9a8c\\u6a21\\u677f") / u("\\u5b9e\\u9a8c\\u6a21\\u677f.docx")
    if sibling_template.exists():
        return sibling_template
    return Path(__file__).resolve().parent.parent / "assets" / "default-template" / u("\\u5b9e\\u9a8c\\u6a21\\u677f.docx")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown_path")
    parser.add_argument("--output")
    parser.add_argument("--template")
    parser.add_argument("--title")
    args = parser.parse_args()

    markdown_path = Path(args.markdown_path)
    markdown_text = markdown_path.read_text(encoding="utf-8")
    output_path = Path(args.output) if args.output else markdown_path.with_suffix(".docx")
    template_path = resolve_template(markdown_path, args.template)
    experiment_title = args.title or guess_experiment_title(markdown_text, markdown_path)

    document = Document(template_path)
    configure_document(document)
    prepare_cover_from_template(document, experiment_title)
    build_body(document, markdown_text)
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
